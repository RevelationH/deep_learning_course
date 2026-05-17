from __future__ import annotations

import json
from datetime import date
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT_DIR = Path(__file__).resolve().parents[1]
ARTIFACT_DIR = ROOT_DIR / "deep_learning_rag" / "artifacts_full_course"
OUTPUT_DIR = Path(__file__).resolve().parent

OUTPUT_DOCX = OUTPUT_DIR / f"report_platform_official_tightened_{date.today():%Y%m%d}.docx"
OUTPUT_MD = OUTPUT_DIR / f"report_platform_official_tightened_{date.today():%Y%m%d}.md"

CHAT_HOME_IMG = ROOT_DIR / "manual_checks" / "20260417_busy_prompt_final" / "busycap07_clean_busy.png"
CHAT_DETAIL_IMG = ROOT_DIR / "manual_checks" / "20260417_busy_prompt_final" / "busycap08_clean_busy.png"
ARCH_IMG = OUTPUT_DIR / "学生端与平台服务架构图.png"
TEACHER_IMG = OUTPUT_DIR / "教师侧内容生产流程图.png"

TITLE_FONT = "方正小标宋简体"
HEADING_FONT = "黑体"
BODY_FONT = "仿宋_GB2312"
TABLE_FONT = "宋体"


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def load_metrics() -> dict[str, Any]:
    inventory = load_json(ARTIFACT_DIR / "inventory.json")
    knowledge_points = load_json(ARTIFACT_DIR / "knowledge_points.json")
    return {
        "inventory": inventory,
        "knowledge_points": knowledge_points,
    }


def set_run_font(run, *, font_name: str, size: float, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_style_font(style, *, font_name: str, size: float, bold: bool = False) -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)


def configure_section(section) -> None:
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


def add_paragraph(
    doc: Document,
    text: str,
    *,
    first_line_indent_cm: float | None = 0.74,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    if first_line_indent_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, font_name=BODY_FONT, size=12)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, font_name=HEADING_FONT, size=16 if level == 1 else 14, bold=True)


def add_numbered_lines(doc: Document, items: Iterable[str]) -> None:
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"{idx}、{item}")
        set_run_font(run, font_name=BODY_FONT, size=12)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 70, start: int = 80, bottom: int = 70, end: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        shade_cell(cell, "F2F2F2")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, font_name=TABLE_FONT, size=10.5, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(value))
            set_run_font(run, font_name=TABLE_FONT, size=10.5)

    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def add_figure(doc: Document, caption: str, path: Path, *, width_inches: float = 6.2) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_inches))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = c.add_run(caption)
    set_run_font(run, font_name=TABLE_FONT, size=10.5, bold=True)


def add_manual_toc(doc: Document) -> None:
    add_paragraph(doc, "本目录按照正文结构编排，页码以正文排版结果为准。", first_line_indent_cm=None, align=WD_ALIGN_PARAGRAPH.LEFT)
    items = [
        ("一、平台概述", "3"),
        ("二、平台功能介绍与展示", "4"),
        ("三、项目完成情况", "6"),
        ("四、运行情况与总体成效", "7"),
        ("五、课程知识组织情况", "8"),
        ("六、教师侧建设内容", "9"),
        ("七、学生侧建设内容", "10"),
        ("八、系统设计与运行机制", "11"),
        ("九、学习报告与教学反馈", "12"),
        ("十、总结", "13"),
        ("附录 A 课程材料统计", "14"),
        ("附录 B 课程来源文件统计", "15"),
        ("附录 C 平台功能清单", "16"),
        ("附录 D 平台界面截图", "17"),
    ]
    for title, page in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
        run1 = p.add_run(title)
        set_run_font(run1, font_name=BODY_FONT, size=12)
        run2 = p.add_run("\t")
        set_run_font(run2, font_name=BODY_FONT, size=12)
        run3 = p.add_run(page)
        set_run_font(run3, font_name=BODY_FONT, size=12)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def write_markdown() -> None:
    lines = [
        "# 深度学习课程大模型平台项目汇报书",
        "",
        "本文件对应 Word 正式汇报稿，正文顺序为平台概述、平台功能介绍与展示、项目完成情况、运行情况与总体成效、课程知识组织情况、教师侧建设内容、学生侧建设内容、系统设计与运行机制、学习报告与教学反馈、总结。",
    ]
    OUTPUT_MD.write_text("\n".join(lines), encoding="utf-8")


def build_report(metrics: dict[str, Any]) -> None:
    inventory = metrics["inventory"]
    knowledge_points = metrics["knowledge_points"]

    doc = Document()
    configure_section(doc.sections[0])

    set_style_font(doc.styles["Normal"], font_name=BODY_FONT, size=12)
    set_style_font(doc.styles["Heading 1"], font_name=HEADING_FONT, size=16, bold=True)
    set_style_font(doc.styles["Heading 2"], font_name=HEADING_FONT, size=14, bold=True)
    set_style_font(doc.styles["Heading 3"], font_name=HEADING_FONT, size=12, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("深度学习课程大模型平台项目汇报书")
    set_run_font(r, font_name=TITLE_FONT, size=24, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目建设完成情况汇报材料")
    set_run_font(r, font_name=BODY_FONT, size=13, bold=True)

    add_table(
        doc,
        ["项目要素", "内容"],
        [
            ["项目名称", "深度学习课程大模型平台"],
            ["材料属性", "项目汇报书"],
            ["建设对象", "深度学习课程"],
            ["编制日期", f"{date.today().year}年{date.today().month:02d}月{date.today().day:02d}日"],
        ],
        [4.5, 11.2],
    )

    add_page_break(doc)
    add_heading(doc, "目录")
    add_manual_toc(doc)
    add_page_break(doc)

    add_heading(doc, "一、平台概述")
    add_paragraph(
        doc,
        "深度学习课程大模型平台是面向课程教学组织、学生学习支持和教学反馈分析建设的课程智能服务平台。平台以课程资料整合为基础，以课程知识组织为主线，以问答、练习、学习报告和历史对话等功能为主要载体，形成了服务课程教学全过程的综合性平台形态。",
    )
    add_paragraph(
        doc,
        "该平台不同于仅承担资料存放或单次交互功能的系统，而是将课程内容组织、学习支持和学习反馈纳入统一平台框架，使课程资料能够由静态载体转化为可检索、可调用、可回看、可持续更新的课程服务资源，较好适应课程教学运行和学生学习使用的实际需要。",
    )

    add_heading(doc, "二、平台功能介绍与展示")
    add_paragraph(
        doc,
        "平台当前已形成面向学生使用的统一入口，主要功能包括课程问答、练习测评、学习报告、历史对话和来源回看。学生可在同一平台内围绕课程内容提出问题、完成练习、查看学习反馈，并根据系统提供的来源定位返回相应课程材料继续复习。",
    )
    add_numbered_lines(
        doc,
        [
            "课程问答。围绕课程概念、模型、训练方法和案例内容进行问答交互，帮助学生理解课程知识。",
            "练习测评。围绕课程知识点组织练习题目，支持学生进行阶段性训练和答案核对。",
            "学习报告。根据学生练习和作答情况形成反馈内容，提示优先复习方向。",
            "历史对话。保存学生既有问答记录，支持延续式学习与回看。",
            "来源回看。将问答与练习内容回连到课程材料相应位置，便于学生定位原始内容。",
        ],
    )
    add_figure(doc, "图 1 平台问答首页界面", CHAT_HOME_IMG)
    add_figure(doc, "图 2 平台问答内容与历史记录界面", CHAT_DETAIL_IMG)
    add_paragraph(
        doc,
        "从平台界面展示情况看，学生能够在统一入口内完成课程问答、练习和学习反馈查看。平台界面设计强调学习过程的连续性、资料回看能力以及功能入口的一致性，能够较好支撑课程学习过程中的连续使用需求。",
    )
    add_figure(doc, "图 3 学生端与平台服务架构图", ARCH_IMG)
    add_figure(doc, "图 4 教师侧内容生产与知识服务流程图", TEACHER_IMG)

    add_heading(doc, "三、项目完成情况")
    add_paragraph(
        doc,
        "本项目已完成深度学习课程智能教学平台的主体建设任务，初步形成了覆盖课程资料整理、课程知识组织、学生学习服务、学习过程记录和教学反馈分析的完整平台体系。平台建设始终围绕课程教学实际需求推进，以提升课程资源利用效率和学生学习支持能力为主要目标。",
    )
    add_paragraph(
        doc,
        "项目完成后，课程资料已实现统一归集，课程知识已实现结构化组织，学生侧已形成问答、练习和学习反馈三类核心服务，平台运行已具备持续更新、持续使用和持续扩展的基本条件。",
    )

    add_heading(doc, "四、运行情况与总体成效")
    add_paragraph(
        doc,
        "目前，平台已经形成“课程资料进入平台、课程知识组织成型、学生使用平台学习、学习结果回流分析”的完整运行链条。课程资料不再以分散文件方式存在，而是被组织为可检索、可定位、可复用的课程知识资源；学生学习活动也不再停留在单次问答或单次练习，而是能够在同一平台内形成连续的学习过程记录。",
    )
    add_paragraph(
        doc,
        "从运行情况看，平台已经能够稳定提供课程问答、知识点练习、学习报告和历史对话等服务；从建设结果看，平台已经完成课程知识库和题库的成型工作，并能够随着课程资料的新增或调整继续更新相关内容。",
    )
    add_paragraph(
        doc,
        "平台建设的重点不在于若干孤立功能的简单叠加，而在于已经形成一套可实际投入课程教学使用的统一平台。教师能够依托平台完成课程内容组织和更新，学生能够依托平台完成课程学习、练习和复习，教学活动中的主要资源和主要流程已在同一体系内实现贯通。",
    )

    add_heading(doc, "五、课程知识组织情况")
    add_paragraph(
        doc,
        "平台已围绕深度学习课程的主要教学内容完成知识点整理工作。知识点设置以课程教学内容为依据，以便于学生理解、检索和复习为原则，能够支撑课程问答、题目组织和学习分析三项基础服务。相关知识点并非简单罗列章节标题，而是按照课程学习中实际需要掌握的内容进行整理和归并。",
    )
    kp_rows = [[kp["name"], kp["description"]] for kp in knowledge_points]
    add_table(doc, ["知识点", "内容说明"], kp_rows, [5.2, 10.5])
    add_paragraph(
        doc,
        "在此基础上，平台能够将课程问答中的问题、练习中的题目以及学习报告中的分析内容统一映射到相应知识点，避免学生在不同模块之间面对不同口径、不同表述和不同来源的内容。",
    )

    add_heading(doc, "六、教师侧建设内容")
    add_paragraph(
        doc,
        "教师侧建设内容主要体现在课程内容整理、课程知识沉淀和教学资源发布三个层面。当前阶段，教师侧未单独设置前台操作页面，而是通过后台流程完成课程内容接入和知识服务发布，其主要作用在于将原始教学资料转化为可直接服务学生学习的课程资源。",
    )
    add_numbered_lines(
        doc,
        [
            "课程资料整理。将课件、教材和补充资料进行统一整理，形成平台的课程内容基础。",
            "课程知识组织。围绕课程主要教学内容整理知识点，并保持知识点与课程材料之间的稳定对应关系。",
            "题库形成与修订。围绕课程知识点组织练习题目，并对题目内容、答案和解析进行校核与调整。",
            "教学资源发布。将整理后的课程知识、题库和来源定位能力统一提供给学生端使用。",
            "后续增量更新。课程资料新增或调整后，可继续在现有平台基础上补充更新，不需要重新搭建整套系统。",
        ],
    )

    add_heading(doc, "七、学生侧建设内容")
    add_paragraph(
        doc,
        "学生侧是平台的主要使用界面，围绕课程学习过程中的提问、练习、复习和回看需求展开。平台已在学生端形成统一入口，学生登录后即可进入问答、练习、学习报告和历史对话等功能页面，在连续的学习环境中完成课程学习相关活动。",
    )
    add_numbered_lines(
        doc,
        [
            "课程问答。学生可围绕课程概念、模型方法、训练机制和课程案例进行提问，并获得面向课程内容的回答。",
            "练习测评。学生可围绕知识点完成练习，查看标准答案、解析内容和相应课程来源。",
            "学习报告。平台可根据学生作答情况形成复习建议，帮助学生识别薄弱知识点和优先复习内容。",
            "历史对话。平台保留学生历史问答记录，便于学生延续前序讨论和持续学习。",
            "来源回看。平台支持学生直接回到课程材料对应位置，减少学习过程中“知道结论但找不到出处”的情况。",
        ],
    )

    add_heading(doc, "八、系统设计与运行机制")
    add_paragraph(
        doc,
        "平台整体按照课程知识组织、业务服务组织和学生使用组织三层关系建设。课程资料进入平台后，并不是直接原样展示，而是先完成课程知识整理和服务组织，再以学生能够直接使用的问答、练习和学习报告形式对外提供。",
    )
    add_numbered_lines(
        doc,
        [
            "课程知识层负责承载课程资料、知识点、题目和讲义来源等基础内容。",
            "服务组织层负责完成课程问答、题目调用、来源定位、学习记录和反馈分析等业务处理。",
            "学生使用层负责统一呈现问答、练习、学习报告和历史对话等界面内容。",
            "数据支撑层负责保存用户信息、学习记录、对话记录和相关结果，保证平台能够持续运行和持续累积学习数据。",
            "整个平台按照统一课程知识体系运行，保证不同功能之间的内容口径保持一致。",
        ],
    )

    add_heading(doc, "九、学习报告与教学反馈")
    add_paragraph(
        doc,
        "学习报告功能用于承接学生学习过程中的结果汇总和后续建议，是平台由内容提供走向学习分析支持的重要环节。该功能能够将学生练习情况、知识点掌握情况和后续复习方向联系起来，形成较为完整的学习反馈机制。",
    )
    add_numbered_lines(
        doc,
        [
            "学习报告能够汇总学生练习结果，并识别需要优先复习的知识点。",
            "学习报告能够提示学生回看相应课程内容，提高复习的针对性。",
            "学习报告也为教师观察学生学习情况提供了可用依据，有助于后续教学调整。",
        ],
    )

    add_heading(doc, "十、总结")
    add_paragraph(
        doc,
        "本项目已完成深度学习课程智能教学平台的主体建设任务，形成了以课程资料为基础、以课程知识为核心、以学生学习服务为主要应用场景的整体平台体系。平台既能够承接课程资料整理和课程知识组织工作，也能够直接服务学生日常学习、练习和复习活动。",
    )
    add_paragraph(
        doc,
        "从项目完成情况看，平台已经具备投入课程教学辅助使用的基本条件。后续可在现有基础上继续优化题库质量、完善学习反馈机制和拓展教师侧使用方式，使平台在课程建设和教学应用中发挥更稳定、更持续的作用。",
    )

    doc.add_page_break()
    add_heading(doc, "附录 A 课程材料统计")
    add_table(
        doc,
        ["周次", "解析单元数"],
        [[week, str(count)] for week, count in inventory["weeks"].items()],
        [4.0, 4.0],
    )

    add_heading(doc, "附录 B 课程来源文件统计")
    add_table(
        doc,
        ["来源文件", "解析单元数"],
        [[name, str(count)] for name, count in inventory["sources"].items()],
        [10.5, 4.0],
    )

    add_heading(doc, "附录 C 平台功能清单")
    add_table(
        doc,
        ["类别", "模块", "主要作用"],
        [
            ["学生端", "登录注册", "完成身份认证并绑定个人学习记录"],
            ["学生端", "课程问答", "围绕课程内容进行连续提问与追问"],
            ["学生端", "练习测评", "按知识点完成选择题训练并查看解析"],
            ["学生端", "学习报告", "查看学习表现、薄弱环节与复习建议"],
            ["数据层", "会话与作答记录", "保存历史对话、答题过程和学习结果"],
            ["知识层", "课程知识库", "保存课程知识点、题库和讲义来源映射"],
            ["运行层", "异步任务与队列", "支撑并发访问和后台任务处理"],
        ],
        [3.2, 4.4, 7.2],
    )

    add_heading(doc, "附录 D 平台界面截图说明")
    add_paragraph(
        doc,
        "本次汇报书正文已纳入平台问答界面和平台总体结构图。知识点练习界面与学习报告界面截图将根据后续正式演示场景补充纳入成稿版本，以保证展示内容与实际运行状态保持一致。",
    )

    doc.save(OUTPUT_DOCX)


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    metrics = load_metrics()
    write_markdown()
    build_report(metrics)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
