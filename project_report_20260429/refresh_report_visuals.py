from __future__ import annotations

import json
import math
import shutil
import tempfile
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from PIL import Image, ImageDraw, ImageFont, ImageOps


ROOT_DIR = Path(r"D:\digital_human\deep_learning")
REPORT_DIR = ROOT_DIR / "project_report_20260429"
ARTIFACT_DIR = ROOT_DIR / "deep_learning_rag" / "artifacts_full_course"
SOURCE_DOCX = REPORT_DIR / "report_platform_official_tightened_20260505.docx"
OUTPUT_DOCX = REPORT_DIR / "report_platform_official_tightened_visuals_20260505.docx"

CHAT_HOME_IMG = ROOT_DIR / "manual_checks" / "20260417_busy_prompt_final" / "busycap07_clean_busy.png"
CHAT_DETAIL_IMG = ROOT_DIR / "manual_checks" / "20260417_busy_prompt_final" / "busycap08_clean_busy.png"
FLOW_STUDENT_IMG = REPORT_DIR / "学生端与平台服务架构图.png"
FLOW_TEACHER_IMG = REPORT_DIR / "教师侧内容生产流程图.png"

CHAT_COMBINED_IMG = REPORT_DIR / "report_chat_combined.png"
QUIZ_RENDER_IMG = REPORT_DIR / "report_quiz_formal.png"
LEARNING_REPORT_RENDER_IMG = REPORT_DIR / "report_learning_report_formal.png"
FLOW_COMBINED_IMG = REPORT_DIR / "report_flow_combined.png"

FONT_UI = Path(r"C:\Windows\Fonts\msyh.ttc")
FONT_UI_BOLD = Path(r"C:\Windows\Fonts\msyhbd.ttc")
FONT_BODY = Path(r"C:\Windows\Fonts\simsun.ttc")

BG = "#f7f7f8"
PANEL = "#ffffff"
PANEL_SOFT = "#fafafa"
INK = "#111827"
MUTED = "#6b7280"
LINE = "#e5e7eb"
LINE_STRONG = "#d1d5db"
BRAND = "#10a37f"
BRAND_SOFT = "#e7f8f4"
ACCENT = "#34d399"
WARN = "#f59e0b"


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def font(path: Path, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(path), size=size)


def draw_round_rect(draw: ImageDraw.ImageDraw, box, radius: int, fill: str, outline: str | None = None, width: int = 1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def wrap_text(text: str, width: int) -> list[str]:
    rows = []
    for raw in str(text).splitlines():
        rows.extend(textwrap_wrap_cjk(raw, width))
    return rows or [""]


def textwrap_wrap_cjk(text: str, width: int) -> list[str]:
    import textwrap

    if len(text) <= width:
        return [text]
    return textwrap.wrap(text, width=width, break_long_words=True, break_on_hyphens=False)


def fit_image(im: Image.Image, target_size: tuple[int, int], bg_color: str = BG) -> Image.Image:
    target_w, target_h = target_size
    canvas = Image.new("RGB", target_size, bg_color)
    fitted = ImageOps.contain(im.convert("RGB"), target_size)
    x = (target_w - fitted.width) // 2
    y = (target_h - fitted.height) // 2
    canvas.paste(fitted, (x, y))
    return canvas


def build_chat_combined() -> Path:
    left = Image.open(CHAT_HOME_IMG).convert("RGB")
    right = Image.open(CHAT_DETAIL_IMG).convert("RGB")
    card_w = 1080
    card_h = 780
    pad = 44
    header_h = 96
    bottom_note_h = 92
    canvas = Image.new("RGB", (pad * 3 + card_w * 2, pad * 2 + header_h + card_h + bottom_note_h), BG)
    draw = ImageDraw.Draw(canvas)

    title_font = font(FONT_UI_BOLD, 34)
    sub_font = font(FONT_UI, 18)
    label_font = font(FONT_UI_BOLD, 20)
    note_font = font(FONT_UI, 17)

    draw.text((pad, 24), "平台问答界面展示", font=title_font, fill=INK)
    draw.text((pad, 66), "左图展示统一入口与历史会话区，右图展示课程问答与来源回看效果。", font=sub_font, fill=MUTED)

    boxes = [
        (pad, pad + header_h, pad + card_w, pad + header_h + card_h),
        (pad * 2 + card_w, pad + header_h, pad * 2 + card_w * 2, pad + header_h + card_h),
    ]
    labels = ["问答首页", "问答与历史记录"]
    images = [left, right]
    for idx, (box, label, source) in enumerate(zip(boxes, labels, images)):
        draw_round_rect(draw, box, 28, PANEL, outline=LINE, width=2)
        x1, y1, x2, y2 = box
        draw_round_rect(draw, (x1 + 20, y1 + 18, x1 + 130, y1 + 56), 14, BRAND_SOFT, outline=None)
        draw.text((x1 + 36, y1 + 26), label, font=label_font, fill=BRAND)
        fitted = fit_image(source, (x2 - x1 - 40, y2 - y1 - 96), PANEL_SOFT)
        canvas.paste(fitted, (x1 + 20, y1 + 74))

    note = (
        "界面采用统一侧边导航，将课程问答、练习测评与学习报告纳入同一学习入口。"
        "学生可以在同一会话内持续提问，并根据历史记录延续前序讨论。"
    )
    draw.text((pad, pad * 2 + header_h + card_h + 18), note, font=note_font, fill=INK)
    canvas.save(CHAT_COMBINED_IMG)
    return CHAT_COMBINED_IMG


def select_quiz_question():
    questions = load_json(ARTIFACT_DIR / "questions.json")
    target_id = "kp-卷积神经网络基础-q5"
    for question in questions:
        if question.get("question_id") == target_id:
            return question
    raise RuntimeError("Target quiz question not found.")


def build_quiz_render() -> Path:
    question = select_quiz_question()
    source_image = Image.open(ARTIFACT_DIR / question["image_path"]).convert("RGB")

    canvas = Image.new("RGB", (1440, 1200), BG)
    draw = ImageDraw.Draw(canvas)
    title_font = font(FONT_UI_BOLD, 30)
    h2_font = font(FONT_UI_BOLD, 26)
    body_font = font(FONT_UI, 18)
    small_font = font(FONT_UI, 15)
    bold_font = font(FONT_UI_BOLD, 18)

    draw_round_rect(draw, (34, 24, 1406, 1170), 30, PANEL, outline=LINE, width=2)
    draw.text((64, 52), "练习测评界面展示", font=title_font, fill=INK)
    draw.text((64, 94), "示例知识点：卷积神经网络基础", font=body_font, fill=MUTED)

    draw_round_rect(draw, (64, 140, 516, 1080), 24, PANEL_SOFT, outline=LINE)
    draw.text((92, 170), "题目配图", font=bold_font, fill=INK)
    draw.text((92, 205), "讲义原图用于支撑题目判断，不是单纯装饰图片。", font=small_font, fill=MUTED)
    image_box = (92, 250, 488, 1030)
    fitted = fit_image(source_image, (image_box[2] - image_box[0], image_box[3] - image_box[1]), "#f3f4f6")
    canvas.paste(fitted, (image_box[0], image_box[1]))

    draw_round_rect(draw, (548, 140, 1376, 1080), 24, "#ffffff", outline=LINE)
    draw.text((580, 174), "题目与选项", font=h2_font, fill=INK)

    q_rows = wrap_text(question["question"], 24)
    y = 224
    for row in q_rows:
        draw.text((580, y), row, font=body_font, fill=INK)
        y += 31

    y += 12
    option_h = 92
    for idx, option in enumerate(question["options"]):
        selected = option.startswith(f"{question['correct_option']}.")
        box = (580, y, 1340, y + option_h)
        draw_round_rect(draw, box, 18, BRAND_SOFT if selected else PANEL_SOFT, outline=LINE_STRONG)
        label = option[:2]
        content = option[3:].strip()
        draw.text((606, y + 18), label, font=bold_font, fill=BRAND if selected else INK)
        rows = wrap_text(content, 26)
        base_y = y + 18
        for j, row in enumerate(rows[:2]):
            draw.text((648, base_y + j * 28), row, font=body_font, fill=INK)
        if selected:
            draw.text((1220, y + 18), "正确答案", font=small_font, fill=BRAND)
        y += option_h + 14

    draw_round_rect(draw, (580, 888, 1340, 1040), 18, "#f0fdf4", outline="#bbf7d0")
    draw.text((606, 912), "答案说明", font=bold_font, fill="#166534")
    exp_rows = wrap_text(question["explanation"], 28)
    for idx, row in enumerate(exp_rows[:3]):
        draw.text((606, 948 + idx * 26), row, font=body_font, fill="#166534")

    draw.text((580, 1060), "该界面支持按知识点进入、查看图文题、提交答案并回看讲义来源。", font=small_font, fill=MUTED)
    canvas.save(QUIZ_RENDER_IMG)
    return QUIZ_RENDER_IMG


def build_learning_report_render() -> Path:
    kps = load_json(ARTIFACT_DIR / "knowledge_points.json")
    focus = [
        ("卷积神经网络基础", 10, 7, 70, "建议继续巩固感受野、卷积层与池化层关系。"),
        ("数学基础与机器学习基础", 8, 5, 62, "建议优先回顾误差分解、偏差方差与不可约噪声。"),
        ("注意力机制与Transformer", 6, 5, 83, "整体表现较稳，可通过少量练习维持熟练度。"),
    ]
    canvas = Image.new("RGB", (1440, 1200), BG)
    draw = ImageDraw.Draw(canvas)

    title_font = font(FONT_UI_BOLD, 30)
    h2_font = font(FONT_UI_BOLD, 24)
    body_font = font(FONT_UI, 18)
    small_font = font(FONT_UI, 15)
    stat_font = font(FONT_UI_BOLD, 34)

    draw_round_rect(draw, (34, 24, 1406, 1170), 30, PANEL, outline=LINE, width=2)
    draw.text((64, 50), "学习报告界面展示", font=title_font, fill=INK)
    draw.text((64, 92), "基于学生作答记录生成复习优先级、知识点表现与后续建议。", font=body_font, fill=MUTED)

    stat_boxes = [
        (64, 146, 320, 286, "累计作答", "24"),
        (346, 146, 602, 286, "综合正确率", "71%"),
        (628, 146, 884, 286, "已覆盖知识点", "8"),
        (910, 146, 1166, 286, "优势知识点", "3"),
        (1192, 146, 1348, 286, "待优先回顾", "2"),
    ]
    for x1, y1, x2, y2, label, value in stat_boxes:
        draw_round_rect(draw, (x1, y1, x2, y2), 20, PANEL_SOFT, outline=LINE)
        draw.text((x1 + 22, y1 + 22), label, font=small_font, fill=MUTED)
        draw.text((x1 + 22, y1 + 68), value, font=stat_font, fill=INK)

    left = (64, 330, 700, 1094)
    right = (740, 330, 1348, 1094)
    draw_round_rect(draw, left, 24, PANEL_SOFT, outline=LINE)
    draw_round_rect(draw, right, 24, PANEL_SOFT, outline=LINE)

    draw.text((92, 360), "知识点表现概览", font=h2_font, fill=INK)
    y = 414
    for name, answered, correct, accuracy, note in focus:
        draw_round_rect(draw, (92, y, 672, y + 172), 18, "#ffffff", outline=LINE)
        draw.text((116, y + 18), name, font=font(FONT_UI_BOLD, 20), fill=INK)
        meta = f"已作答 {answered} 题    答对 {correct} 题    正确率 {accuracy}%"
        draw.text((116, y + 52), meta, font=small_font, fill=MUTED)
        track = (116, y + 86, 620, y + 102)
        draw_round_rect(draw, track, 8, "#e5e7eb")
        fill_w = int((track[2] - track[0]) * accuracy / 100)
        draw_round_rect(draw, (track[0], track[1], track[0] + fill_w, track[3]), 8, BRAND if accuracy >= 75 else WARN)
        for idx, row in enumerate(wrap_text(note, 28)[:2]):
            draw.text((116, y + 118 + idx * 24), row, font=body_font, fill=INK)
        y += 192

    draw.text((768, 360), "优先复习建议", font=h2_font, fill=INK)
    recs = [
        "优先回顾“数学基础与机器学习基础”。当前薄弱点集中在误差分解与不可约噪声的理解上。",
        "继续巩固“卷积神经网络基础”。建议结合讲义中的感受野示意图重新做同主题题目。",
        "“注意力机制与Transformer”整体表现较稳，可保留少量复习频次维持熟练度。",
    ]
    y = 418
    for rec in recs:
        draw_round_rect(draw, (768, y, 1318, y + 118), 18, "#ffffff", outline=LINE)
        rows = wrap_text(rec, 27)
        for idx, row in enumerate(rows[:3]):
            draw.text((794, y + 20 + idx * 28), row, font=body_font, fill=INK)
        y += 138

    draw.text((768, 860), "推荐下一步练习", font=h2_font, fill=INK)
    followups = [
        "从卷积网络基础题组中继续完成 5 题，重点关注局部纹理与上下文关系。",
        "返回课程讲义中关于误差分解的页面，先复习公式含义，再完成对应知识点练习。",
    ]
    y = 910
    for rec in followups:
        draw_round_rect(draw, (768, y, 1318, y + 96), 18, "#ffffff", outline=LINE)
        rows = wrap_text(rec, 27)
        for idx, row in enumerate(rows[:2]):
            draw.text((794, y + 18 + idx * 28), row, font=body_font, fill=INK)
        y += 116

    canvas.save(LEARNING_REPORT_RENDER_IMG)
    return LEARNING_REPORT_RENDER_IMG


def build_flow_combined() -> Path:
    left = fit_image(Image.open(FLOW_STUDENT_IMG), (860, 980), PANEL)
    right = fit_image(Image.open(FLOW_TEACHER_IMG), (860, 980), PANEL)
    canvas = Image.new("RGB", (1800, 1120), BG)
    draw = ImageDraw.Draw(canvas)
    title_font = font(FONT_UI_BOLD, 34)
    sub_font = font(FONT_UI, 18)
    label_font = font(FONT_UI_BOLD, 20)
    draw.text((48, 26), "平台整体架构与内容生产流程", font=title_font, fill=INK)
    draw.text((48, 68), "左图展示学生端服务结构，右图展示教师侧课程资料进入知识服务平台的流程。", font=sub_font, fill=MUTED)
    boxes = [(48, 112, 896, 1088), (904, 112, 1752, 1088)]
    labels = ["学生端服务结构", "教师侧内容生产流程"]
    for (x1, y1, x2, y2), label, img in zip(boxes, labels, [left, right]):
        draw_round_rect(draw, (x1, y1, x2, y2), 28, PANEL, outline=LINE, width=2)
        draw_round_rect(draw, (x1 + 24, y1 + 18, x1 + 220, y1 + 58), 14, BRAND_SOFT)
        draw.text((x1 + 42, y1 + 26), label, font=label_font, fill=BRAND)
        canvas.paste(img, (x1 + 16, y1 + 76))
    canvas.save(FLOW_COMBINED_IMG)
    return FLOW_COMBINED_IMG


def replace_media(docx_in: Path, docx_out: Path, replacements: dict[str, Path]):
    tmpdir = Path(tempfile.mkdtemp(prefix="report_refresh_"))
    try:
        with ZipFile(docx_in) as zin:
            zin.extractall(tmpdir)
        for media_name, src_path in replacements.items():
            target = tmpdir / "word" / "media" / media_name
            shutil.copyfile(src_path, target)
        with ZipFile(docx_out, "w") as zout:
            for file in sorted(tmpdir.rglob("*")):
                if file.is_file():
                    zout.write(file, file.relative_to(tmpdir).as_posix())
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def refresh_captions(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    replace_map = {
        "图 1 平台问答首页界面": "图 1 平台问答界面（合并展示）",
        "图 2 平台问答内容与历史记录界面": "图 2 知识点练习界面",
        "图 3 学生端与平台服务架构图": "图 3 学习报告界面",
        "图 4 教师侧内容生产与知识服务流程图": "图 4 平台整体架构与内容生产流程",
    }
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text in replace_map:
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = replace_map[text]
            else:
                paragraph.add_run(replace_map[text])
    doc.save(str(docx_path))


def main():
    chat = build_chat_combined()
    quiz = build_quiz_render()
    learning_report = build_learning_report_render()
    flow = build_flow_combined()
    replace_media(
        SOURCE_DOCX,
        OUTPUT_DOCX,
        {
            "image1.png": chat,
            "image2.png": quiz,
            "image3.png": learning_report,
            "image4.png": flow,
        },
    )
    refresh_captions(OUTPUT_DOCX)
    print(str(OUTPUT_DOCX))


if __name__ == "__main__":
    main()
