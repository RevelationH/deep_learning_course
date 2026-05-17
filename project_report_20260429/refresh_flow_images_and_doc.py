from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document


BASE = Path(r"D:\digital_human\deep_learning\project_report_20260429")
DOCX_PATH = BASE / "report_platform_official_tightened_visuals_splitflows_fixed_20260506_images_updated_v3.docx"
OUTPUT_PATH = BASE / "report_platform_official_tightened_visuals_splitflows_fixed_20260506_images_updated_v4.docx"

EDGE_CANDIDATES = [
    Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"),
    Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe"),
    Path(r"C:\Program Files\Mozilla Firefox\firefox.exe"),
]

SVG_TO_PNG = [
    (
        BASE / "学生端与平台服务架构图.svg",
        BASE / "学生端与平台服务架构图.png",
        "word/media/image4.png",
        (1800, 1160),
    ),
    (
        BASE / "教师侧内容生产流程图.svg",
        BASE / "教师侧内容生产流程图.png",
        "word/media/image5.png",
        (1800, 1120),
    ),
]


def find_browser() -> Path:
    for candidate in EDGE_CANDIDATES:
        if candidate.exists():
            return candidate
    raise FileNotFoundError("No supported browser renderer found.")


def render_svg(browser: Path, svg_path: Path, png_path: Path, size: tuple[int, int]) -> None:
    png_path.parent.mkdir(parents=True, exist_ok=True)
    uri = svg_path.resolve().as_uri()
    cmd = [
        str(browser),
        "--headless",
        "--disable-gpu",
        f"--window-size={size[0]},{size[1]}",
        f"--screenshot={png_path}",
        uri,
    ]
    subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    if not png_path.exists() or png_path.stat().st_size == 0:
        raise RuntimeError(f"Failed to render {svg_path.name}")


def replace_docx_media(docx_path: Path, output_path: Path, replacements: list[tuple[Path, str]]) -> Path:
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_root = Path(temp_dir)
        with ZipFile(docx_path, "r") as zin:
            zin.extractall(temp_root)
        for src, inner_name in replacements:
            target = temp_root / inner_name
            target.parent.mkdir(parents=True, exist_ok=True)
            shutil.copyfile(src, target)
        with ZipFile(output_path, "w", ZIP_DEFLATED) as zout:
            for path in sorted(temp_root.rglob("*")):
                if path.is_file():
                    zout.write(path, path.relative_to(temp_root).as_posix())
    return output_path


def update_appendices(docx_path: Path) -> None:
    doc = Document(docx_path)

    appendix_a = {
        "Week1": "深度学习概述与发展背景",
        "Week2": "数学基础与机器学习基础",
        "Week3": "数学基础补充与卷积网络引入",
        "Week4": "卷积神经网络基础结构",
        "Week5": "典型卷积网络与视觉应用",
        "Week6": "循环神经网络与 LSTM",
        "Week7": "注意力机制与 Transformer",
        "Week8": "GPT、BERT、视觉 Transformer 与 GAN 基础",
        "Week9": "GAN 进阶与对抗生成训练",
        "Week10": "深度生成模型与扩散模型",
        "Week11": "正则化、泛化与集成学习",
        "Week12": "优化方法与训练策略",
        "Week13": "开发平台与课程综合案例",
    }
    appendix_b = {
        "0.1-DeepLearning-Introduction-C1.pdf": "深度学习概述、发展背景与表示学习思想",
        "0.2-DeepLearning-Foundations-C2.pdf": "数学基础、机器学习基本概念与模型训练基础",
        "0.2-DeepLearning-Foundations-C3.pdf": "线性代数、概率统计与课程所需数学工具",
        "0.3-DeepLearing-CNN-C3.pdf": "卷积神经网络引入、卷积与池化基本思想",
        "0.3-DeepLearing-CNN-C4.pdf": "CNN 基本结构、感受野、参数共享与特征提取",
        "0.3-DeepLearing-CNN-C5.pdf": "LeNet、AlexNet、VGG、ResNet 等典型卷积网络",
        "0.4-DeepLearing-RNN-C6.pdf": "RNN、BPTT、LSTM 与序列建模",
        "0.5-Transformer-C7.pdf": "注意力机制、自注意力与 Transformer 基础",
        "0.5-Transformer-C8.pdf": "GPT、BERT、Swin Transformer 等模型",
        "0.6-GAN-C8.pdf": "GAN 基本结构、生成器与判别器对抗训练",
        "0.6-GAN-C9.pdf": "条件 GAN、训练稳定性与改进策略",
        "0.6-GenerativeModel-C10.pdf": "自编码器、RBM、DBN、扩散模型等生成模型",
        "0.7-RegularizationOptimization-C11.pdf": "正则化方法、泛化误差与训练稳定性",
        "0.7-RegularizationOptimization-C12.pdf": "优化方法、归一化、训练技巧与调参",
        "0.8-DeepLearning-tools-C13_modified.pdf": "TensorFlow、PyTorch、Keras 等开发平台",
        "DeepForest-C11.pdf": "决策树、随机森林与 DeepForest",
    }

    if len(doc.tables) >= 5:
        week_table = doc.tables[2]
        week_table.cell(0, 0).text = "周次"
        week_table.cell(0, 1).text = "主要教学主题"
        for row in week_table.rows[1:]:
            week = row.cells[0].text.strip()
            if week in appendix_a:
                row.cells[1].text = appendix_a[week]

        file_table = doc.tables[3]
        file_table.cell(0, 0).text = "来源文件"
        file_table.cell(0, 1).text = "主要内容"
        for row in file_table.rows[1:]:
            file_name = row.cells[0].text.strip()
            if file_name in appendix_b:
                row.cells[1].text = appendix_b[file_name]

        feature_table = doc.tables[4]
        existing_rows = [[cell.text.strip() for cell in row.cells] for row in feature_table.rows]
        teacher_rows = [
            ["教师侧", "课程资料接入", "导入课件、教材与讲义等课程资料"],
            ["教师侧", "知识库与题库生成", "完成课程知识组织、题目生成与来源映射"],
            ["教师侧", "审核与持续更新", "对知识点、题目和课程内容进行抽检修订并保持更新"],
        ]
        existing_first_two = {(row[0], row[1]) for row in existing_rows if len(row) >= 2}
        for row_data in teacher_rows:
            key = (row_data[0], row_data[1])
            if key in existing_first_two:
                continue
            row = feature_table.add_row()
            for idx, value in enumerate(row_data):
                row.cells[idx].text = value

    doc.save(docx_path)


def main() -> None:
    browser = find_browser()
    replacements: list[tuple[Path, str]] = []
    for svg_path, png_path, inner_name, size in SVG_TO_PNG:
        render_svg(browser, svg_path, png_path, size)
        replacements.append((png_path, inner_name))
    output_path = replace_docx_media(DOCX_PATH, OUTPUT_PATH, replacements)
    update_appendices(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
