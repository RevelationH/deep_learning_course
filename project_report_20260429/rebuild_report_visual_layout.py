from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw


REPORT_DIR = Path(__file__).resolve().parent
SOURCE_DOC = REPORT_DIR / "report_platform_official_tightened_visuals_explained_20260506.docx"
OUTPUT_DOC = REPORT_DIR / "report_platform_official_tightened_visuals_splitflows_fixed_20260506.docx"

CHAT_IMG = Path(r"D:\digital_human\deep_learning\manual_checks\20260417_busy_prompt_final\busycap01_result.png")
QUIZ_IMG = REPORT_DIR / "report_quiz_formal.png"
LEARNING_IMG = REPORT_DIR / "report_learning_report_formal.png"
STUDENT_FLOW_IMG = REPORT_DIR / "学生端与平台服务架构图.png"
TEACHER_FLOW_RAW = REPORT_DIR / "教师侧内容生产流程图.png"
TEACHER_FLOW_FIXED = REPORT_DIR / "teacher_flow_fixed_20260506.png"

BODY_FONT = "仿宋_GB2312"
HEADING_FONT = "黑体"
TABLE_FONT = "宋体"


FIGURES = [
    (
        "图 1 课程问答界面",
        "图1展示了课程问答功能的实际使用界面。学生可围绕课程概念、模型结构、训练方法和案例内容开展连续提问，系统在回答过程中同步提供课程来源定位信息，以支持学生在理解结论的同时回看原始课程材料。",
        CHAT_IMG,
        14.8,
    ),
    (
        "图 2 知识点练习界面",
        "图2展示了知识点练习界面。系统按照课程知识点组织题目，并同步提供题干、配图、选项、答案解析和来源定位信息，便于学生在完成作答后及时核对掌握情况，进一步围绕薄弱环节开展针对性复习。",
        QUIZ_IMG,
        14.8,
    ),
    (
        "图 3 学习报告界面",
        "图3展示了学习报告界面。系统基于学生作答记录汇总知识点表现，识别需要优先复习的内容，并给出后续练习与材料回看建议，以支持学生形成阶段性的学习总结与复习安排。",
        LEARNING_IMG,
        14.8,
    ),
    (
        "图 4 学生端服务架构图",
        "图4展示了学生端服务架构。图中依次说明了学生访问入口、统一业务编排层、课程问答服务、练习测评服务、学习报告服务以及底层知识资产与运行支撑之间的关系，用于说明学生侧各项功能的协同方式。",
        STUDENT_FLOW_IMG,
        14.8,
    ),
    (
        "图 5 教师端内容生产流程图",
        "图5展示了教师端内容生产流程。图中说明课程资料接入、内容清洗、知识组织、知识库与题库形成、审核修订、发布到学生端以及持续更新等环节之间的衔接关系，用于说明教师侧后台内容生产与发布机制。",
        TEACHER_FLOW_FIXED,
        14.8,
    ),
]


def set_run_font(run, font_name: str, size: float, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def clear_paragraph(paragraph: Paragraph) -> None:
    while paragraph.runs:
        paragraph._p.remove(paragraph.runs[0]._r)


def replace_paragraph_text(paragraph: Paragraph, text: str, font_name: str, size: float, bold: bool = False) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, font_name, size, bold)


def style_body_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = Pt(28)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def style_caption_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = Pt(22)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def style_image_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(0)


def next_paragraph(paragraph: Paragraph) -> Paragraph | None:
    node = paragraph._p.getnext()
    while node is not None:
        if node.tag == qn("w:p"):
            return Paragraph(node, paragraph._parent)
        node = node.getnext()
    return None


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def remove_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def clean_directory_page(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "目录":
            style_caption_paragraph(paragraph)
            replace_paragraph_text(paragraph, "目录", HEADING_FONT, 16, True)
        elif text == "本目录按照正文结构编排，页码以正文排版结果为准。":
            remove_paragraph(paragraph)
        elif "\t" in text and "、" in text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.left_indent = Cm(0.74)
            paragraph.paragraph_format.line_spacing = Pt(28)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, BODY_FONT, 12)


def set_picture(paragraph: Paragraph, image_path: Path, width_cm: float) -> None:
    clear_paragraph(paragraph)
    style_image_paragraph(paragraph)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))


def fix_teacher_flow() -> Path:
    im = Image.open(TEACHER_FLOW_RAW).convert("RGB")
    draw = ImageDraw.Draw(im)
    color = (0, 101, 155)
    # reconnect the missing downward arrow under “持续更新”
    draw.line((1350, 690, 1350, 752), fill=color, width=6)
    draw.polygon([(1350, 760), (1342, 744), (1358, 744)], fill=color)
    im.save(TEACHER_FLOW_FIXED)
    return TEACHER_FLOW_FIXED


def rebuild_figure_section(doc: Document) -> None:
    start_idx = None
    end_idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == "图 1 平台问答界面（合并展示）":
            start_idx = i - 1 if i > 0 and doc.paragraphs[i - 1]._p.xpath(".//w:drawing") else i
        if text == "三、项目完成情况":
            end_idx = i
            break

    if start_idx is None or end_idx is None or start_idx >= end_idx:
        raise RuntimeError("Could not locate figure section.")

    anchor = doc.paragraphs[start_idx - 1] if start_idx > 0 else doc.paragraphs[0]
    for paragraph in list(doc.paragraphs[start_idx:end_idx]):
        remove_paragraph(paragraph)

    current = anchor
    for caption, explanation, image_path, width_cm in FIGURES:
        img_para = insert_paragraph_after(current)
        set_picture(img_para, image_path, width_cm)

        caption_para = insert_paragraph_after(img_para)
        style_caption_paragraph(caption_para)
        replace_paragraph_text(caption_para, caption, TABLE_FONT, 10.5, True)

        explanation_para = insert_paragraph_after(caption_para)
        style_body_paragraph(explanation_para)
        replace_paragraph_text(explanation_para, explanation, BODY_FONT, 12)

        spacer = insert_paragraph_after(explanation_para)
        spacer.alignment = WD_ALIGN_PARAGRAPH.LEFT
        spacer.paragraph_format.first_line_indent = Cm(0)
        spacer.paragraph_format.line_spacing = Pt(10)
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)
        replace_paragraph_text(spacer, "", BODY_FONT, 12)
        current = spacer


def refresh_appendix_note(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "附录 D 平台界面截图说明":
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(2)
            replace_paragraph_text(paragraph, text, HEADING_FONT, 16, True)
            nxt = next_paragraph(paragraph)
            if nxt is not None:
                style_body_paragraph(nxt)
                replace_paragraph_text(
                    nxt,
                    "附录 D 所列界面截图已在正文中纳入课程问答、知识点练习、学习报告、学生端服务架构图及教师端内容生产流程图等内容，用于说明平台面向学生的实际使用界面以及教师侧后台内容生产的整体组织方式。",
                    BODY_FONT,
                    12,
                )
            break


def main() -> None:
    fix_teacher_flow()
    doc = Document(str(SOURCE_DOC))
    clean_directory_page(doc)
    rebuild_figure_section(doc)
    refresh_appendix_note(doc)
    doc.save(str(OUTPUT_DOC))
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
