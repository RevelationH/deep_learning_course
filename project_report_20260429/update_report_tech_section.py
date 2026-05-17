from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


BASE_DOC = Path(r"D:\digital_human\deep_learning\project_report_20260429\report_platform_official.docx")
BACKUP_DOC = BASE_DOC.with_name("report_platform_official_before_tech_section_backup.docx")


TOC_INSERT_TITLE = "五、平台采用的关键技术"
TECH_SECTION_TITLE = "五、平台采用的关键技术"

TECH_PARAGRAPHS = [
    "为支撑课程问答、知识点练习、学习报告、来源回看和历史对话等功能稳定运行，平台在底层采用了检索增强生成、语义向量检索、结构化文档解析、知识点映射、异步任务编排和多后端数据持久化等关键技术。相关技术并非孤立部署，而是围绕“课程资料进入、课程知识组织、学生交互服务生成、学习结果沉淀反馈”这一主链路形成协同运行机制。",
    "1、检索增强生成（Retrieval-Augmented Generation，RAG）技术。平台将回答过程拆分为“语义检索+答案生成”两段式流程，先在课程知识库内完成相关内容召回，再据此组织回答，从而使问答结果保持在课程语境范围内，并支持答案与原始课程材料之间的来源回连。",
    "2、语义向量表示与近似最近邻检索技术。平台采用 HuggingFace Embeddings 框架和 sentence-transformers/all-MiniLM-L6-v2 嵌入模型，将课程文本表示为向量特征，并基于 FAISS 构建向量索引，实现面向语义相似度的内容召回，避免仅依赖关键词匹配导致的检索偏差。",
    "3、课程材料解析与结构化切片技术。平台采用 PyMuPDF（fitz）与 pypdf 对课程 PDF 资料进行分页解析，并结合固定窗口切分、重叠上下文保留和元数据标注等方法，将原始材料转化为可检索、可定位、可复用的课程知识片段，为问答、题目解析和报告生成提供统一内容基础。",
    "4、知识点映射与来源定位技术。平台将知识点体系作为统一语义主线，将问答内容、练习题目、答案解析和学习报告结果映射到相应知识点，并同步保留页码、文件名、周次和标题等来源元数据，从而实现“结论可解释、出处可回看、内容可追溯”的课程知识服务方式。",
    "5、异步任务编排与后台 Worker 技术。针对问答生成和学习报告生成等耗时任务，平台采用异步队列与后台 Worker 执行机制，将前台请求与后台处理过程解耦，并通过任务状态跟踪、队列位置管理和结果回收机制提升高并发条件下的运行稳定性。",
    "6、数据持久化与运行支撑技术。平台采用 PostgreSQL / Firebase 双后端适配方案保存用户信息、历史对话、作答记录和学习报告结果，并结合 Redis JSON 缓存、会话持久化和快照保存机制，提高高频访问场景下的数据读写效率与运行连续性。",
    "上述技术共同构成了平台的底层支撑体系，使课程资料能够转化为可直接服务教学与学习的知识资产，并使学生端的问答、练习和学习反馈功能在同一技术框架下保持内容一致、来源一致和结果可追溯。",
]


BODY_HEADING_UPDATES = {
    "五、课程知识组织情况": "六、课程知识组织情况",
    "六、教师侧建设内容": "七、教师侧建设内容",
    "七、学生侧建设内容": "八、学生侧建设内容",
    "八、系统设计与运行机制": "九、系统设计与运行机制",
    "九、学习报告与教学反馈": "十、学习报告与教学反馈",
    "十、总结": "十一、总结",
}


TOC_TEXTS_TO_SHIFT = [
    "五、课程知识组织情况",
    "六、教师侧建设内容",
    "七、学生侧建设内容",
    "八、系统设计与运行机制",
    "九、学习报告与教学反馈",
    "十、总结",
    "附录 A 课程材料统计",
    "附录 B 课程来源文件统计",
    "附录 C 平台功能清单",
]


def find_paragraph(doc: Document, exact_text: str) -> Paragraph:
    for para in doc.paragraphs:
        if para.text.strip() == exact_text:
            return para
    raise ValueError(f"Paragraph not found: {exact_text}")


def make_empty_paragraph_from_template(template: Paragraph, anchor: Paragraph) -> Paragraph:
    new_p = deepcopy(template._p)
    for child in list(new_p):
        if not child.tag.endswith("}pPr"):
            new_p.remove(child)
    anchor._p.addprevious(new_p)
    return Paragraph(new_p, anchor._parent)


def insert_before(reference: Paragraph, text: str, template: Paragraph) -> Paragraph:
    new_para = make_empty_paragraph_from_template(template, reference)
    new_para.add_run(text)
    return new_para


def parse_toc_line(text: str) -> tuple[str, int] | None:
    match = re.match(r"^(.*)\t(\d+)$", text.strip())
    if not match:
        return None
    return match.group(1), int(match.group(2))


def update_toc(doc: Document) -> None:
    old_five = None
    five_page = 8
    for para in doc.paragraphs:
        parsed = parse_toc_line(para.text)
        if not parsed:
            continue
        title, page = parsed
        if title == "五、课程知识组织情况":
            old_five = para
            five_page = page
            break
    if old_five is None:
        raise ValueError("Unable to locate TOC line for section five.")

    insert_before(old_five, f"{TOC_INSERT_TITLE}\t{five_page}", old_five)

    for para in doc.paragraphs:
        text = para.text.strip()
        parsed = parse_toc_line(text)
        if not parsed:
            continue
        title, page = parsed
        if title == "五、课程知识组织情况":
            para.text = f"六、课程知识组织情况\t{page + 1}"
        elif title == "六、教师侧建设内容":
            para.text = f"七、教师侧建设内容\t{page + 1}"
        elif title == "七、学生侧建设内容":
            para.text = f"八、学生侧建设内容\t{page + 1}"
        elif title == "八、系统设计与运行机制":
            para.text = f"九、系统设计与运行机制\t{page + 1}"
        elif title == "九、学习报告与教学反馈":
            para.text = f"十、学习报告与教学反馈\t{page + 1}"
        elif title == "十、总结":
            para.text = f"十一、总结\t{page + 1}"
        elif title == "附录 A 课程材料统计":
            para.text = f"附录 A 课程材料统计\t{page + 1}"
        elif title == "附录 B 课程来源文件统计":
            para.text = f"附录 B 课程来源文件统计\t{page + 1}"
        elif title == "附录 C 平台功能清单":
            para.text = f"附录 C 平台功能清单\t{page + 1}"


def renumber_body_headings(doc: Document) -> None:
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in BODY_HEADING_UPDATES:
            para.text = BODY_HEADING_UPDATES[text]


def insert_tech_section(doc: Document) -> None:
    next_heading = find_paragraph(doc, "六、课程知识组织情况")
    heading_template = next_heading
    body_template = None
    for para in doc.paragraphs:
        if para.text.strip().startswith("在此基础上"):
            body_template = para
            break
    if body_template is None:
        body_template = find_paragraph(doc, "在此基础上，平台能够将课程问答中的问题、练习中的题目以及学习报告中的分析内容统一映射到相应知识点，避免学生在不同模块之间面对不同口径、不同表述和不同来源的内容。")

    insert_before(next_heading, TECH_SECTION_TITLE, heading_template)
    for text in TECH_PARAGRAPHS:
        insert_before(next_heading, text, body_template)


def main() -> None:
    if not BASE_DOC.exists():
        raise FileNotFoundError(BASE_DOC)

    shutil.copy2(BASE_DOC, BACKUP_DOC)

    doc = Document(BASE_DOC)
    update_toc(doc)
    renumber_body_headings(doc)
    insert_tech_section(doc)
    doc.save(BASE_DOC)

    print(f"Updated: {BASE_DOC}")
    print(f"Backup:  {BACKUP_DOC}")


if __name__ == "__main__":
    main()
