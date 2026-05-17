from __future__ import annotations

import html
import json
import math
import shutil
import sys
from datetime import date
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT_DIR = Path(__file__).resolve().parents[1]
ARTIFACT_DIR = ROOT_DIR / "deep_learning_rag" / "artifacts_full_course"
OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_DOCX = OUTPUT_DIR / f"深度学习课程大模型平台项目汇报书_修订版_{date.today():%Y%m%d}.docx"
OUTPUT_MD = OUTPUT_DIR / f"深度学习课程大模型平台项目汇报书_修订版_{date.today():%Y%m%d}.md"
FIG_TEACHER = OUTPUT_DIR / "教师侧内容生产流程图.png"
FIG_ARCH = OUTPUT_DIR / "学生端与平台服务架构图.png"
FIG_TEACHER_SVG = OUTPUT_DIR / "教师侧内容生产流程图.svg"
FIG_ARCH_SVG = OUTPUT_DIR / "学生端与平台服务架构图.svg"
FIG_UI_CHAT_HOME = OUTPUT_DIR / "平台问答首页截图.png"
FIG_UI_CHAT_DETAIL = OUTPUT_DIR / "平台问答内容截图.png"

TITLE_FONT_NAME = "方正小标宋简体"
BODY_FONT_NAME = "仿宋_GB2312"
HEADING_FONT_NAME = "黑体"
TABLE_FONT_NAME = "宋体"


def load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def load_metrics() -> Dict[str, Any]:
    build_meta = load_json(ARTIFACT_DIR / "build_meta.json")
    inventory = load_json(ARTIFACT_DIR / "inventory.json")
    kps = load_json(ARTIFACT_DIR / "knowledge_points.json")
    questions = load_json(ARTIFACT_DIR / "questions.json")
    image_question_count = sum(1 for item in questions if item.get("image_path"))
    kp_rows = []
    for item in kps:
        kp_rows.append(
            {
                "name": item.get("name", ""),
                "weeks": "、".join(item.get("weeks", [])),
                "description": item.get("description", ""),
                "source_files": "；".join(item.get("source_files", [])),
            }
        )
    return {
        "build_meta": build_meta,
        "inventory": inventory,
        "knowledge_points": kps,
        "questions": questions,
        "image_question_count": image_question_count,
        "kp_rows": kp_rows,
    }


def _find_font() -> str:
    candidates = [
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\simkai.ttf"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\msyh.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return str(path)
    return ""


FONT_PATH = _find_font()


def get_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    if FONT_PATH:
        return ImageFont.truetype(FONT_PATH, size=size)
    return ImageFont.load_default()


def ensure_text(text: str) -> str:
    return str(text).replace("\n", " ").strip()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> List[str]:
    if not text:
        return [""]
    lines: List[str] = []
    for paragraph in str(text).splitlines():
        current = ""
        for ch in paragraph:
            test = current + ch
            bbox = draw.textbbox((0, 0), test, font=font)
            width = bbox[2] - bbox[0]
            if width <= max_width or not current:
                current = test
            else:
                lines.append(current)
                current = ch
        lines.append(current or "")
    return lines or [""]


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: Tuple[int, int, int] | None = None,
    font_name: str = BODY_FONT_NAME,
):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = RGBColor(*(color or (0, 0, 0)))


def set_style_font(style, *, font_name: str = BODY_FONT_NAME, size: float | None = None, bold: bool | None = None):
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    style: str | None = None,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_cm: float | None = 0.74,
    font_name: str = BODY_FONT_NAME,
):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if first_line_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_cm)
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=12, font_name=font_name)
    return p


def add_numbered_lines(doc: Document, items: Iterable[str]) -> None:
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"{idx}、{ensure_text(item)}")
        set_run_font(run, size=12, font_name=BODY_FONT_NAME)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 14, 3: 12}[level], bold=True, font_name=HEADING_FONT_NAME)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 70, start: int = 80, bottom: int = 70, end: int = 80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_simple_table(doc: Document, headers: List[str], rows: List[List[str]], col_widths_cm: List[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        hdr_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(hdr_cells[index])
        shade_cell(hdr_cells[index], "F2F2F2")
        hdr_cells[index].text = ""
        p = hdr_cells[index].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=10.5, bold=True, font_name=TABLE_FONT_NAME)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[index])
            cells[index].text = ""
            p = cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(ensure_text(value))
            set_run_font(run, size=10.5, font_name=TABLE_FONT_NAME)
    if col_widths_cm:
        for row in table.rows:
            for idx, width in enumerate(col_widths_cm):
                row.cells[idx].width = Cm(width)


def add_figure(doc: Document, title: str, image_path: Path, width_inches: float = 6.8) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(4)
    run = caption.add_run(title)
    set_run_font(run, size=10.5, bold=True, font_name=TABLE_FONT_NAME)


def set_section_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)


def set_section_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


def add_landscape_figure(doc: Document, title: str, image_path: Path, *, width_inches: float = 9.0) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_landscape(section)
    add_figure(doc, title, image_path, width_inches=width_inches)
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_portrait(section)


def add_directory_entry(doc: Document, title: str, page: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.first_line_indent = Cm(0)
    tabs = p.paragraph_format.tab_stops
    tabs.clear_all()
    tabs.add_tab_stop(Cm(15.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
    r1 = p.add_run(title)
    set_run_font(r1, size=12, font_name=BODY_FONT_NAME)
    r2 = p.add_run("\t")
    set_run_font(r2, size=12, font_name=BODY_FONT_NAME)
    r3 = p.add_run(page)
    set_run_font(r3, size=12, font_name=BODY_FONT_NAME)


def add_manual_toc(doc: Document) -> None:
    add_paragraph(
        doc,
        "本目录按照正文结构编排，页码以正文排版结果为准。",
        first_line_cm=None,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    directory_items = [
        ("一、平台概述", "3"),
        ("二、平台功能介绍与展示", "4"),
        ("三、项目完成情况", "6"),
        ("四、运行情况与总体成效", "7"),
        ("五、课程知识组织情况", "8"),
        ("六、教师侧建设内容", "9"),
        ("七、学生侧建设内容", "10"),
        ("八、系统设计与运行机制", "11"),
        ("九、学习报告与教学反馈", "12"),
        ("十、总结", "13"),
        ("附录 A 课程材料统计", "14"),
        ("附录 B 课程来源文件统计", "15"),
        ("附录 C 平台功能清单", "16"),
        ("附录 D 平台界面截图", "17"),
    ]
    for title, page in directory_items:
        add_directory_entry(doc, title, page)


def make_box(x1: int, y1: int, x2: int, y2: int, title: str, lines: List[str], title_fill: Tuple[int, int, int]) -> Dict[str, Any]:
    return {
        "x1": x1,
        "y1": y1,
        "x2": x2,
        "y2": y2,
        "title": title,
        "lines": lines,
        "title_fill": title_fill,
        "fill": (255, 255, 255),
        "outline": (52, 86, 112),
    }


def color_hex(rgb: Tuple[int, int, int]) -> str:
    return "#{:02x}{:02x}{:02x}".format(*rgb)


def svg_text_block(x: int, y: int, lines: List[str], *, font_size: int, fill: str, font_weight: str = "400") -> str:
    if not lines:
        return ""
    parts = [
        f'<text x="{x}" y="{y}" font-family="Microsoft YaHei, SimHei, sans-serif" font-size="{font_size}" font-weight="{font_weight}" fill="{fill}">'
    ]
    for idx, line in enumerate(lines):
        escaped = html.escape(line)
        if idx == 0:
            parts.append(f'<tspan x="{x}" dy="0">{escaped}</tspan>')
        else:
            parts.append(f'<tspan x="{x}" dy="{int(font_size * 1.45)}">{escaped}</tspan>')
    parts.append("</text>")
    return "".join(parts)


def svg_box(spec: Dict[str, Any]) -> str:
    x1, y1, x2, y2 = spec["x1"], spec["y1"], spec["x2"], spec["y2"]
    width = x2 - x1
    height = y2 - y1
    title_fill = color_hex(spec["title_fill"])
    outline = color_hex(spec["outline"])
    fill = color_hex(spec["fill"])
    title_svg = svg_text_block(x1 + 18, y1 + 34, [spec["title"]], font_size=24, fill="#ffffff", font_weight="700")
    body_svg = svg_text_block(x1 + 18, y1 + 88, spec["lines"], font_size=19, fill="#202020")
    return (
        f'<g>'
        f'<rect x="{x1}" y="{y1}" width="{width}" height="{height}" rx="24" ry="24" fill="{fill}" stroke="{outline}" stroke-width="3"/>'
        f'<path d="M {x1+24} {y1} H {x2-24} A 24 24 0 0 1 {x2} {y1+24} V {y1+52} H {x1} V {y1+24} A 24 24 0 0 1 {x1+24} {y1} Z" fill="{title_fill}"/>'
        f'<rect x="{x1}" y="{y1+28}" width="{width}" height="24" fill="{title_fill}"/>'
        f"{title_svg}{body_svg}</g>"
    )


def svg_arrow(start: Tuple[int, int], end: Tuple[int, int], *, color: str = "#0c6aa6", width: int = 5) -> str:
    x1, y1 = start
    x2, y2 = end
    angle = math.atan2(y2 - y1, x2 - x1)
    head_len = 15
    head_angle = math.pi / 7
    p1 = (x2 - head_len * math.cos(angle - head_angle), y2 - head_len * math.sin(angle - head_angle))
    p2 = (x2 - head_len * math.cos(angle + head_angle), y2 - head_len * math.sin(angle + head_angle))
    return (
        f'<g>'
        f'<line x1="{x1}" y1="{y1}" x2="{x2}" y2="{y2}" stroke="{color}" stroke-width="{width}" stroke-linecap="round"/>'
        f'<polygon points="{x2},{y2} {p1[0]:.1f},{p1[1]:.1f} {p2[0]:.1f},{p2[1]:.1f}" fill="{color}"/>'
        f'</g>'
    )


def write_svg_figure(
    path: Path,
    *,
    width: int,
    height: int,
    title: str,
    subtitle: str,
    boxes: List[Dict[str, Any]],
    arrows: List[Tuple[Tuple[int, int], Tuple[int, int]]],
) -> None:
    elements = [
        f'<rect x="0" y="0" width="{width}" height="{height}" fill="#f8fbfe"/>',
        svg_text_block(60, 72, [title], font_size=36, fill="#16324c", font_weight="700"),
        svg_text_block(60, 118, [subtitle], font_size=22, fill="#556676"),
    ]
    for box in boxes:
        elements.append(svg_box(box))
    for start, end in arrows:
        elements.append(svg_arrow(start, end))
    path.write_text(
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">{"".join(elements)}</svg>',
        encoding="utf-8",
    )


def render_diagram_png(
    path: Path,
    *,
    width: int,
    height: int,
    title: str,
    subtitle: str,
    boxes: List[Dict[str, Any]],
    arrows: List[Tuple[Tuple[int, int], Tuple[int, int]]],
) -> None:
    image = Image.new("RGB", (width, height), (248, 251, 254))
    draw = ImageDraw.Draw(image)
    draw.text((58, 34), title, font=get_font(34), fill=(22, 44, 69))
    draw.text((60, 84), subtitle, font=get_font(21), fill=(85, 102, 118))
    for box in boxes:
        x1, y1, x2, y2 = box["x1"], box["y1"], box["x2"], box["y2"]
        title_fill = box["title_fill"]
        fill = box["fill"]
        outline = box["outline"]
        draw.rounded_rectangle((x1, y1, x2, y2), radius=22, fill=fill, outline=outline, width=3)
        header_h = 50
        draw.rounded_rectangle((x1, y1, x2, y1 + header_h), radius=22, fill=title_fill, outline=title_fill)
        draw.rectangle((x1, y1 + header_h - 22, x2, y1 + header_h), fill=title_fill, outline=title_fill)
        draw.text((x1 + 18, y1 + 10), box["title"], font=get_font(24), fill=(255, 255, 255))
        cursor_y = y1 + header_h + 14
        for line in box["lines"]:
            for wrapped in wrap_text(draw, line, get_font(19), max(100, x2 - x1 - 36)):
                draw.text((x1 + 18, cursor_y), wrapped, font=get_font(19), fill=(32, 32, 32))
                cursor_y += 29
    for start, end in arrows:
        draw.line([start, end], fill=(0, 102, 153), width=6)
        angle = math.atan2(end[1] - start[1], end[0] - start[0])
        head_len = 16
        head_angle = math.pi / 8
        p1 = (
            end[0] - head_len * math.cos(angle - head_angle),
            end[1] - head_len * math.sin(angle - head_angle),
        )
        p2 = (
            end[0] - head_len * math.cos(angle + head_angle),
            end[1] - head_len * math.sin(angle + head_angle),
        )
        draw.polygon([end, p1, p2], fill=(0, 102, 153))
    image.save(path)


def create_teacher_figure(metrics: Dict[str, Any]) -> None:
    width, height = 1750, 980
    boxes = [
        make_box(70, 170, 370, 350, "资料接入", ["课件、教材、参考讲义", "统一导入课程资源池"], (35, 90, 140)),
        make_box(430, 170, 760, 350, "内容清洗", ["抽取文本、标题、页码与图像位置", "清理噪声并保留来源关系"], (35, 90, 140)),
        make_box(820, 170, 1150, 350, "知识组织", ["形成知识点、关键词与页码映射", "构建课程知识骨架"], (35, 90, 140)),
        make_box(1210, 170, 1680, 350, "知识库与题库生成", [f"形成 {metrics['build_meta']['dense_index']['count']} 个检索单元", f"形成 {metrics['build_meta']['question_count']} 道练习题"], (35, 90, 140)),
        make_box(150, 470, 600, 690, "审核与修订", ["人工抽检知识点与题目质量", "调整来源定位与内容边界"], (47, 116, 102)),
        make_box(690, 470, 1090, 690, "发布到学生端", ["同步问答、练习和学习报告", "支撑前端可视化页面展示"], (47, 116, 102)),
        make_box(1180, 470, 1620, 690, "持续更新", ["新增课件后重新构建知识库", "保持课程内容与教学进度一致"], (118, 58, 130)),
        make_box(500, 760, 1260, 900, "管理要点", ["教师侧以内容生产和审核为主，不强调独立入口界面，而是通过后台流程完成课程资源沉淀、题库发布与知识库更新。"], (118, 58, 130)),
    ]
    arrows = [
        ((370, 260), (430, 260)),
        ((760, 260), (820, 260)),
        ((1150, 260), (1210, 260)),
        ((980, 350), (980, 470)),
        ((600, 580), (690, 580)),
        ((1090, 580), (1180, 580)),
        ((1350, 690), (1350, 760)),
    ]
    title = "教师侧内容生产与知识服务流程图"
    subtitle = "从课程资料接入、内容处理、知识组织到发布和更新，展示教师侧的完整工作链路。"
    render_diagram_png(FIG_TEACHER, width=width, height=height, title=title, subtitle=subtitle, boxes=boxes, arrows=arrows)
    write_svg_figure(FIG_TEACHER_SVG, width=width, height=height, title=title, subtitle=subtitle, boxes=boxes, arrows=arrows)


def create_architecture_figure(metrics: Dict[str, Any]) -> None:
    width, height = 1800, 1120
    boxes = [
        make_box(120, 160, 1680, 300, "学生访问层", ["登录/注册、课程问答、练习测评、学习报告、历史对话"], (35, 90, 140)),
        make_box(120, 380, 1680, 530, "统一门厅与业务编排层", ["负责身份识别、会话管理、页面路由、请求汇聚和结果组织"], (35, 90, 140)),
        make_box(120, 610, 560, 850, "智能问答服务", ["结合课程检索与上下文进行回答", "优先提供课程来源与讲义定位"], (47, 116, 102)),
        make_box(680, 610, 1120, 850, "练习测评服务", ["按知识点组织题目", "支持图文题与自动评分"], (47, 116, 102)),
        make_box(1240, 610, 1680, 850, "学习报告服务", ["统计答题情况与知识点覆盖", "给出后续复习建议"], (47, 116, 102)),
        make_box(120, 940, 920, 1080, "课程知识资产层", [f"课程 PDF、知识点、题库与讲义图像资源；当前检索单元数为 {metrics['build_meta']['dense_index']['count']}。"], (35, 90, 140)),
        make_box(960, 940, 1680, 1080, "数据与运行支撑层", ["保存用户、会话、答题记录、学习报告快照与任务状态，支撑并发访问与持续运行。"], (118, 58, 130)),
    ]
    arrows = [
        ((900, 300), (900, 380)),
        ((900, 530), (340, 610)),
        ((900, 530), (900, 610)),
        ((900, 530), (1460, 610)),
        ((340, 850), (520, 940)),
        ((900, 850), (900, 940)),
        ((1460, 850), (1320, 940)),
    ]
    title = "学生端与平台服务架构图"
    subtitle = "从访问入口、业务编排、课程知识服务到数据与运行支撑，展示平台的整体协同方式。"
    render_diagram_png(FIG_ARCH, width=width, height=height, title=title, subtitle=subtitle, boxes=boxes, arrows=arrows)
    write_svg_figure(FIG_ARCH_SVG, width=width, height=height, title=title, subtitle=subtitle, boxes=boxes, arrows=arrows)


def create_ui_figures() -> None:
    chat_home = FIG_UI_CHAT_HOME
    chat_detail = FIG_UI_CHAT_DETAIL
    src_home = ROOT_DIR / "manual_checks" / "20260417_busy_prompt_final" / "busycap07_result_round2.png"
    src_detail = ROOT_DIR / "manual_checks" / "20260417_busy_prompt_final" / "busycap09_result_round2.png"
    if src_home.exists():
        shutil.copy2(src_home, chat_home)
    if src_detail.exists():
        shutil.copy2(src_detail, chat_detail)


def add_platform_screenshots(doc: Document) -> None:
    src_home = ROOT_DIR / "manual_checks" / "20260417_busy_prompt_final" / "busycap07_clean_busy.png"
    src_detail = ROOT_DIR / "manual_checks" / "20260417_busy_prompt_final" / "busycap08_clean_busy.png"
    if src_home.exists():
        add_figure(doc, "图 1 平台问答首页界面", src_home, width_inches=6.2)
    if src_detail.exists():
        add_figure(doc, "图 2 平台问答内容与历史记录界面", src_detail, width_inches=6.2)


def add_core_figures(doc: Document) -> None:
    if FIG_ARCH.exists():
        add_figure(doc, "图 3 学生端与平台服务架构图", FIG_ARCH, width_inches=6.2)
    if FIG_TEACHER.exists():
        add_figure(doc, "图 4 教师侧内容生产与知识服务流程图", FIG_TEACHER, width_inches=6.2)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run("目录")
    set_run_font(run, size=16, bold=True, font_name=HEADING_FONT_NAME)
    toc = doc.add_paragraph()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    r = toc.add_run()
    r._r.append(fld_char_begin)
    r._r.append(instr_text)
    r._r.append(fld_char_sep)
    r._r.append(fld_char_end)
    set_run_font(r, size=11, font_name=BODY_FONT_NAME)


def write_markdown(metrics: Dict[str, Any]) -> None:
    build_meta = metrics["build_meta"]
    inventory = metrics["inventory"]
    lines = []
    lines.append("# 深度学习课程大模型平台项目汇报书")
    lines.append("")
    lines.append("## 一、项目完成情况")
    lines.append("")
    lines.append(
        "本项目已完成面向深度学习课程的课程知识服务平台建设，形成了课程资料导入、知识库构建、课程问答、练习测评、学习报告和历史对话管理等核心能力。"
    )
    lines.append("")
    lines.append("## 二、平台建设结果")
    lines.append("")
    lines.append(f"- 课程 PDF 文件数量：{inventory.get('file_count', 0)}")
    lines.append(f"- 原始教学单元数量：{inventory.get('raw_document_count', 0)}")
    lines.append(f"- 知识点数量：{build_meta.get('knowledge_point_count', 0)}")
    lines.append(f"- 题目数量：{build_meta.get('question_count', 0)}")
    lines.append(f"- 带图题数量：{metrics.get('image_question_count', 0)}")
    lines.append(f"- Dense 检索单元数量：{build_meta.get('dense_index', {}).get('count', 0)}")
    lines.append("")
    lines.append("## 三、课程知识点覆盖")
    lines.append("")
    for item in metrics["kp_rows"]:
        lines.append(f"- {item['name']}：{item['description']}")
    lines.append("")
    lines.append("## 四、附录")
    lines.append("")
    lines.append("- 平台截图：问答首页、问答内容页")
    lines.append("- 结构图：教师侧内容生产流程图、学生端与平台服务架构图")
    OUTPUT_MD.write_text("\n".join(lines), encoding="utf-8")


def add_toc_field(doc: Document) -> None:
    p = doc.add_paragraph()
    r = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_char_begin)
    r._r.append(instr_text)
    r._r.append(fld_char_sep)
    r._r.append(fld_char_end)
    set_run_font(r, size=11, font_name=BODY_FONT_NAME)


def build_report(metrics: Dict[str, Any]) -> None:
    inventory = metrics["inventory"]
    kp_rows = metrics["kp_rows"]

    doc = Document()
    section = doc.sections[0]
    set_section_portrait(section)

    set_style_font(doc.styles["Normal"], font_name=BODY_FONT_NAME, size=12)
    set_style_font(doc.styles["Heading 1"], font_name=HEADING_FONT_NAME, size=16, bold=True)
    set_style_font(doc.styles["Heading 2"], font_name=HEADING_FONT_NAME, size=14, bold=True)
    set_style_font(doc.styles["Heading 3"], font_name=HEADING_FONT_NAME, size=12, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    r = title.add_run("????????????????")
    set_run_font(r, size=24, bold=True, font_name=TITLE_FONT_NAME)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("????????????")
    set_run_font(r, size=13, bold=True, font_name=BODY_FONT_NAME)

    meta_table = [
        ["????", "???????????"],
        ["????", "?????"],
        ["????", "??????"],
        ["????", f"{date.today().year}?{date.today().month:02d}?{date.today().day:02d}?"],
    ]
    add_simple_table(doc, ["????", "??"], meta_table, [4.5, 11.2])

    doc.add_page_break()
    add_heading(doc, "??", 1)
    add_manual_toc(doc)

    add_heading(doc, "??????", 1)
    add_paragraph(
        doc,
        "????????????????????????????????????????????????????????????????????????????????????????????????????????????????????????????",
    )
    add_paragraph(
        doc,
        "??????????????????????????????????????????????????????????????????????????????????????????????????????",
    )

    add_heading(doc, "???????????", 1)
    add_paragraph(
        doc,
        "????????????????????????????????????????????????????????????????????????????????????????????????????????????",
    )
    add_numbered_lines(
        doc,
        [
            "??????????????????????????????????????????",
            "????????????????????????????????????",
            "????????????????????????????????",
            "???????????????????????????",
            "????????????????????????????????????",
        ],
    )
    add_platform_screenshots(doc)
    add_paragraph(
        doc,
        "???????????????????????????????????????????????????????????????????????",
    )
    add_core_figures(doc)

    add_heading(doc, "????????", 1)
    add_paragraph(
        doc,
        "????????????????????????????????????????????????????????????????????????????????????????????????????????????",
    )
    add_paragraph(
        doc,
        "?????????????????????????????????????????????????????????????????????????????????",
    )

    add_heading(doc, "???????????", 1)
    add_paragraph(
        doc,
        "???????????????????????????????????????????????????????????????????????????????????????????????????????????????????????????????????????????",
    )
    add_paragraph(
        doc,
        "????????????????????????????????????????????????????????????????????????????????????????????",
    )
    add_paragraph(
        doc,
        "?????????????????????????????????????????????????????????????????????????????????????????????????????????????????",
    )

    add_heading(doc, "??????????", 1)
    add_paragraph(
        doc,
        "????????????????????????????????????????????????????????????????????????????????????????????????????????????????????????????????",
    )
    kp_table = [[row["name"], row["description"]] for row in kp_rows]
    add_simple_table(doc, ["???", "????"], kp_table, [5.2, 10.5])
    add_paragraph(
        doc,
        "????????????????????????????????????????????????????????????????????????????????",
    )

    add_heading(doc, "?????????", 1)
    add_paragraph(
        doc,
        "???????????????????????????????????????????????????????????????????????????????????????????????????????????????????",
    )
    add_numbered_lines(
        doc,
        [
            "?????????????????????????????????????",
            "????????????????????????????????????????????",
            "??????????????????????????????????????????",
            "?????????????????????????????????????",
            "??????????????????????????????????????????????",
        ],
    )

    add_heading(doc, "?????????", 1)
    add_paragraph(
        doc,
        "??????????????????????????????????????????????????????????????????????????????????????????????????????????",
    )
    add_numbered_lines(
        doc,
        [
            "???????????????????????????????????????????????",
            "?????????????????????????????????????",
            "??????????????????????????????????????????",
            "??????????????????????????????????",
            "???????????????????????????????????????????????",
        ],
    )

    add_heading(doc, "???????????", 1)
    add_paragraph(
        doc,
        "??????????????????????????????????????????????????????????????????????????????????????????????????",
    )
    add_numbered_lines(
        doc,
        [
            "???????????????????????????????",
            "???????????????????????????????????????",
            "????????????????????????????????",
            "?????????????????????????????????????????????????",
            "???????????????????????????????????",
        ],
    )

    add_heading(doc, "???????????", 1)
    add_paragraph(
        doc,
        "????????????????????????????????????????????????????????????????????????????????????????????????????",
    )
    add_numbered_lines(
        doc,
        [
            "?????????????????????????????",
            "????????????????????????????",
            "??????????????????????????????????",
        ],
    )

    add_heading(doc, "????", 1)
    add_paragraph(
        doc,
        "??????????????????????????????????????????????????????????????????????????????????????????????????????????????????",
    )
    add_paragraph(
        doc,
        "???????????????????????????????????????????????????????????????????????????????????????????????",
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "?? A ??????", 1)
    add_simple_table(doc, ["??", "?????"], [[week, str(count)] for week, count in inventory.get("weeks", {}).items()], [4.0, 4.0])

    add_heading(doc, "?? B ????????", 1)
    add_simple_table(doc, ["????", "?????"], [[name, str(count)] for name, count in inventory.get("sources", {}).items()], [10.5, 4.0])

    add_heading(doc, "?? C ??????", 1)
    add_simple_table(
        doc,
        ["??", "??", "????"],
        [
            ["???", "????", "???????????????"],
            ["???", "????", "???????????????"],
            ["???", "????", "????????????????"],
            ["???", "????", "????????????????"],
            ["???", "???????", "????????????????"],
            ["???", "?????", "?????????????????"],
            ["???", "???????", "?????????????"],
        ],
        [3.2, 4.4, 7.2],
    )

    add_heading(doc, "?? D ??????", 1)
    add_platform_screenshots(doc)

    doc.save(OUTPUT_DOCX)


def main() -> None:
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    metrics = load_metrics()
    create_teacher_figure(metrics)
    create_architecture_figure(metrics)
    create_ui_figures()
    write_markdown(metrics)
    build_report(metrics)
    print(f"Report written to: {OUTPUT_DOCX}")
    print(f"Markdown written to: {OUTPUT_MD}")
    print(f"Figures written to: {FIG_TEACHER} ; {FIG_ARCH} ; {FIG_TEACHER_SVG} ; {FIG_ARCH_SVG}")
    print(f"UI screenshots written to: {FIG_UI_CHAT_HOME} ; {FIG_UI_CHAT_DETAIL}")


if __name__ == "__main__":
    main()
