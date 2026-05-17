from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOCX_PATH = Path(r"D:\digital_human\deep_learning\project_report_20260429\report_platform_official.docx")
FALLBACK_PATH = DOCX_PATH.with_name("report_platform_official_updated.docx")


def insert_paragraph_after(paragraph: Paragraph, text: str, style_name: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = style_name
    new_para.text = text
    return new_para


def main() -> None:
    doc = Document(DOCX_PATH)
    paragraphs = doc.paragraphs

    toc_insert_after = None
    body_insert_after = None

    for para in paragraphs:
        text = para.text.strip()
        if text.startswith("四、运行情况与总体成效\t") and toc_insert_after is None:
            toc_insert_after = para
        if text == "四、运行情况与总体成效":
            body_insert_after = para

    if toc_insert_after is None or body_insert_after is None:
        raise RuntimeError("Could not locate insertion points.")

    # Update TOC entries after section 4.
    toc_replacements = {
        "五、课程知识组织情况\t8": "六、课程知识组织情况\t9",
        "六、教师侧建设内容\t9": "七、教师侧建设内容\t10",
        "七、学生侧建设内容\t10": "八、学生侧建设内容\t11",
        "八、系统设计与运行机制\t11": "九、系统设计与运行机制\t12",
        "九、学习报告与教学反馈\t12": "十、学习报告与教学反馈\t13",
        "十、总结\t13": "十一、总结\t14",
        "附录 A 课程材料统计\t14": "附录 A 课程材料统计\t15",
        "附录 B 课程来源文件统计\t15": "附录 B 课程来源文件统计\t16",
        "附录 C 平台功能清单\t16": "附录 C 平台功能清单\t17",
    }
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in toc_replacements:
            para.text = toc_replacements[text]

    insert_paragraph_after(toc_insert_after, "五、平台采用的关键专业技术\t8", toc_insert_after.style.name)

    # Renumber subsequent body headings.
    heading_replacements = {
        "五、课程知识组织情况": "六、课程知识组织情况",
        "六、教师侧建设内容": "七、教师侧建设内容",
        "七、学生侧建设内容": "八、学生侧建设内容",
        "八、系统设计与运行机制": "九、系统设计与运行机制",
        "九、学习报告与教学反馈": "十、学习报告与教学反馈",
        "十、总结": "十一、总结",
    }
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in heading_replacements:
            para.text = heading_replacements[text]

    current = body_insert_after
    current = insert_paragraph_after(current, "五、平台采用的关键专业技术", body_insert_after.style.name)
    normal_style = "Normal"

    tech_paragraphs = [
        "本平台在总体实现上采用“课程知识组织 + 检索增强生成 + 学习行为记录 + 反馈分析联动”的技术路线。其核心目标不是将课程材料简单搬运到前端页面，而是将原始教学资料转化为可检索、可定位、可回连、可复用的课程知识服务体系，使平台在问答、练习、学习报告和历史对话等功能之间保持统一的数据基础与统一的知识口径。",
        "1、检索增强生成技术。平台在课程问答环节采用检索增强生成（Retrieval-Augmented Generation，RAG）机制，将课程 PDF、知识点、题目来源和讲义图片等内容组织为可检索知识资源。在学生提问后，系统并非直接进行无约束生成，而是先基于课程知识资源完成语义检索与来源筛选，再将检索结果与当前会话上下文共同送入生成环节，从而提升回答的课程相关性、来源可追溯性和内容稳定性。该机制能够有效降低泛化回答对课程边界的偏离，支持学生在获得结论的同时同步查看相应课程出处。",
        "2、课程材料解析与知识切片技术。平台对课程资料采用分层处理方式，对教学文件中的标题、页码、正文片段、图像位置和知识点归属关系进行解析和重组，将原本连续的讲义内容转化为面向教学服务的知识片段集合。通过这种知识切片与来源映射机制，平台能够在问答中完成页级来源定位，在练习模块中完成题目与知识点绑定，在学习报告中完成薄弱知识点归因，从而保证不同模块所调用的信息来源一致、解释口径一致、回看路径一致。",
        "3、知识点驱动的题库生成与题目组织技术。平台在题库侧采用“知识点主线 + 课程材料约束 + 来源映射保留”的组织方法，将课程重点概念、模型结构、训练方法和案例内容转化为可用于练习的题目资源。题目并非作为孤立文本保存，而是与知识点、答案解析、来源页码和相关配图形成关联关系。基于这一组织方式，系统能够围绕指定知识点发放练习，能够在答题后返回解析与出处，也能够在学习报告阶段按知识点维度统计表现并生成后续复习建议。",
        "4、会话持久化与学习状态管理技术。平台对学生的历史对话、作答记录、学习报告结果和来源回看状态进行统一持久化管理，使课程问答不再是单次输入输出，而是具备可延续、可回看、可分析的学习过程属性。会话状态的保留支持学生围绕同一主题进行追问，作答记录的保留支持后续学习报告生成，学习结果的保留则为阶段性复习、重点内容回顾和持续学习反馈提供数据基础。该机制使平台能够从“即时回答系统”扩展为“过程性学习支持系统”。",
        "5、异步任务与运行支撑技术。考虑到平台同时承载课程问答、练习测评、学习报告和历史记录等多项功能，系统在运行层引入了异步任务处理和后台状态管理机制，用于支撑请求分发、结果写入、状态回收和并发访问控制。对于学习报告生成、历史会话读取、题目结果记录等需要跨模块协同的数据处理任务，平台通过统一的数据支撑层和运行支撑层进行承接，从而保持前端功能响应的稳定性，并为后续课程资料更新、知识库再构建和服务扩展保留可持续运行的技术基础。",
        "综合来看，平台采用的关键专业技术并不是若干孤立算法的堆叠，而是围绕课程教学场景形成的一套集知识组织、语义检索、来源定位、过程记录和学习反馈于一体的技术体系。该体系既保证了课程内容服务的专业性与可追溯性，也为平台后续扩展教师侧管理功能、增强学习分析能力和拓展课程应用范围提供了明确的技术基础。",
    ]

    for text in tech_paragraphs:
        current = insert_paragraph_after(current, text, normal_style)

    try:
        doc.save(DOCX_PATH)
        print(DOCX_PATH)
    except PermissionError:
        doc.save(FALLBACK_PATH)
        print(FALLBACK_PATH)


if __name__ == "__main__":
    main()
