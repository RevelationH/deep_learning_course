from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


DOC_PATH = Path(r"D:\digital_human\deep_learning\project_report_20260429\report_platform_official.docx")
BACKUP_PATH = DOC_PATH.with_name("report_platform_official_before_heading_student_fix_backup.docx")


TARGET_HEADING_INDEXES = [56, 65, 68, 75, 82, 89, 94]

STUDENT_SECTION_TEXTS = {
    76: "学生侧建设内容主要体现为围绕学生个人学习过程形成持续、可回溯、可反馈的使用机制。第二部分已经对平台功能作了总体介绍，本部分侧重说明学生在实际使用过程中能够形成的学习支持方式，以及平台如何围绕个人学习过程提供连续服务。",
    77: "1、形成连续学习链路。学生可在同一账户下持续保留既往提问、作答情况和学习反馈，不必在问答、练习和复习之间反复切换系统或重新建立学习上下文。",
    78: "2、形成面向个人的练习记录。平台能够按知识点保存学生作答结果，便于学生区分已经掌握、尚未掌握和需要重点回看的内容，使练习结果由一次性使用转化为可累积的学习记录。",
    79: "3、形成针对薄弱环节的反馈支持。系统可根据学生作答表现生成学习报告，提示优先复习方向，并将后续练习建议与课程材料回看路径衔接起来，提高复习安排的针对性。",
    80: "4、形成可追溯的来源回看机制。学生在查看问答结果、题目解析和学习建议时，可同步定位相应课程材料位置，减少学习过程中“知道结论但无法回到原始内容”的情况。",
    81: "5、形成持续性的课程互动环境。历史对话保留机制使学生能够在既有讨论基础上继续追问、补充和延展，不必每次从零开始输入背景信息，从而提高课程学习的连续性和使用效率。",
}


def apply_heading_run_format(reference_para, target_para) -> None:
    ref_font_name = None
    if reference_para.runs:
        ref_font_name = reference_para.runs[0].font.name
    if not ref_font_name:
        ref_font_name = "SimHei"

    for run in target_para.runs:
        run.font.name = ref_font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), ref_font_name)
        run._element.rPr.rFonts.set(qn("w:ascii"), ref_font_name)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), ref_font_name)


def main() -> None:
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)

    shutil.copy2(DOC_PATH, BACKUP_PATH)

    doc = Document(DOC_PATH)

    reference_heading = doc.paragraphs[52]
    for idx in TARGET_HEADING_INDEXES:
        apply_heading_run_format(reference_heading, doc.paragraphs[idx])

    for idx, text in STUDENT_SECTION_TEXTS.items():
        doc.paragraphs[idx].text = text

    doc.save(DOC_PATH)
    print(f"Updated: {DOC_PATH}")
    print(f"Backup:  {BACKUP_PATH}")


if __name__ == "__main__":
    main()
