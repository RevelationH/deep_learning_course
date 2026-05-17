from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


REPORT_DIR = Path(__file__).resolve().parent
SOURCE_DOC = REPORT_DIR / "report_platform_official_tightened_visuals_20260505.docx"
OUTPUT_DOC = REPORT_DIR / "report_platform_official_tightened_visuals_explained_20260506.docx"

BODY_FONT = "仿宋_GB2312"
CAPTION_FONT = "宋体"
HEADING_PATTERN = re.compile(r"^(?:[一二三四五六七八九十]+、|附录)")
LEADING_PUNCT_PATTERN = re.compile(r"^[\s。，“”‘’、；：！？.,;:)\]）】》]+")

CAPTION_EXPLANATIONS = {
    "图 1 平台问答界面（合并展示）": (
        "图1展示了平台问答功能的主要使用界面。界面左侧集中呈现功能入口与历史对话，右侧为课程问答交互区域，"
        "学生可围绕课程概念、模型方法和案例内容开展连续提问，并结合来源回看功能进一步定位相关课程材料。"
    ),
    "图 2 知识点练习界面": (
        "图2展示了知识点练习界面。系统按照课程知识点组织选择题练习，并同步提供题干、选项、答案解析和来源定位信息，"
        "便于学生在完成作答后及时核对理解情况，进而围绕薄弱知识点开展针对性复习。"
    ),
    "图 3 学习报告界面": (
        "图3展示了学习报告界面。系统基于学生作答情况汇总学习表现，识别需要优先复习的知识点，并给出后续练习与材料回看建议，"
        "以支持学生形成阶段性的复习安排。"
    ),
    "图 4 平台整体架构与内容生产流程": (
        "图4展示了平台整体架构与内容生产流程。教师侧通过课程资料整理、知识点组织和题库形成完成内容生产，"
        "学生侧通过问答、练习和学习报告等功能开展学习使用，平台后台则负责知识组织、记录保存和服务联动。"
    ),
}


def set_run_font(run, font_name: str = BODY_FONT, size: float = 12, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def style_body_paragraph(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = Pt(28)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)


def replace_paragraph_text(paragraph: Paragraph, text: str, *, font_name: str = BODY_FONT, bold: bool = False) -> None:
    while paragraph.runs:
        paragraph._p.remove(paragraph.runs[0]._r)
    run = paragraph.add_run(text)
    set_run_font(run, font_name=font_name, size=12 if font_name == BODY_FONT else 10.5, bold=bold)


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    style_body_paragraph(new_paragraph)
    run = new_paragraph.add_run(text)
    set_run_font(run, font_name=BODY_FONT, size=12)
    return new_paragraph


def next_paragraph(paragraph: Paragraph) -> Paragraph | None:
    node = paragraph._p.getnext()
    while node is not None:
        if node.tag == qn("w:p"):
            return Paragraph(node, paragraph._parent)
        node = node.getnext()
    return None


def clean_leading_punctuation(text: str) -> str:
    return LEADING_PUNCT_PATTERN.sub("", text)


def normalize_paragraph(paragraph: Paragraph) -> None:
    text = paragraph.text
    cleaned = clean_leading_punctuation(text)
    if cleaned != text:
        replace_paragraph_text(paragraph, cleaned)


def normalize_tables(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    normalize_paragraph(paragraph)


def is_caption(text: str) -> bool:
    return text.strip().startswith("图 ")


def is_heading(text: str) -> bool:
    return bool(HEADING_PATTERN.match(text.strip()))


def ensure_explanation_after_caption(paragraph: Paragraph, explanation: str) -> None:
    nxt = next_paragraph(paragraph)
    if nxt is None:
        insert_paragraph_after(paragraph, explanation)
        return

    if not nxt.text.strip():
        style_body_paragraph(nxt)
        replace_paragraph_text(nxt, explanation)
        return

    text = nxt.text.strip()
    if text.startswith("从平台界面展示情况看"):
        style_body_paragraph(nxt)
        replace_paragraph_text(nxt, explanation)
        return

    if is_caption(text) or is_heading(text):
        insert_paragraph_after(paragraph, explanation)
        return

    style_body_paragraph(nxt)
    replace_paragraph_text(nxt, explanation)


def refresh_appendix_note(doc: Document) -> None:
    old_text = (
        "本次汇报书正文已纳入平台问答界面和平台总体结构图。知识点练习界面与学习报告界面截图将根据后续正式演示场景补充纳入成稿版本，以保证展示内容与实际运行状态保持一致。"
    )
    new_text = (
        "附录所列界面截图与正文展示内容一致，已覆盖平台问答、知识点练习和学习报告等主要学生使用界面，"
        "可用于说明平台的实际运行状态、主要交互方式以及学生侧的连续学习流程。"
    )
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == old_text:
            style_body_paragraph(paragraph)
            replace_paragraph_text(paragraph, new_text)
            break


def main() -> None:
    doc = Document(SOURCE_DOC)

    for paragraph in doc.paragraphs:
        normalize_paragraph(paragraph)

    normalize_tables(doc)

    for paragraph in list(doc.paragraphs):
        caption = paragraph.text.strip()
        explanation = CAPTION_EXPLANATIONS.get(caption)
        if explanation:
            ensure_explanation_after_caption(paragraph, explanation)

    refresh_appendix_note(doc)
    doc.save(OUTPUT_DOC)
    print(OUTPUT_DOC)


if __name__ == "__main__":
    main()
